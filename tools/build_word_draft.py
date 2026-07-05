# -*- coding: utf-8 -*-
from __future__ import annotations

import re
import json
from datetime import date
from pathlib import Path

from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
INDEX = ROOT / "site" / "index.html"
OUT = ROOT / "docs" / "wrtn-product-assignment-draft.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(23, 27, 34)
MUTED = RGBColor(91, 100, 114)
LIGHT_FILL = "F2F4F7"
CALLOUT_FILL = "F4F6F9"
BORDER = "D9DEE8"


def set_rfonts(run, ascii_font="Calibri", east_asia_font="Malgun Gothic"):
    run.font.name = ascii_font
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)
    rfonts.set(qn("w:eastAsia"), east_asia_font)


def set_para_spacing(paragraph, before=0, after=6, line=1.10):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def style_document(doc: Document):
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        st = doc.styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        st.font.size = Pt(size)
        st.font.color.rgb = color
        st.font.bold = True
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.line_spacing = 1.10

    footer = section.footer.paragraphs[0]
    footer.text = "Wrtn AX Product Manager (Ontology) Assignment Draft"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(footer, 0, 0, 1.0)
    for run in footer.runs:
        set_rfonts(run)
        run.font.size = Pt(8.5)
        run.font.color.rgb = MUTED


def add_hyperlink(paragraph, text: str, url: str):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2B50D8")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(color)
    rpr.append(underline)
    new_run.append(rpr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    new_run.append(text_el)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def normalize_href(href: str) -> str:
    if href.startswith("#"):
        return ""
    if href.startswith(("http://", "https://", "file://", "mailto:")):
        return href
    if href.startswith("app.html"):
        return (ROOT / "site" / "app.html").resolve().as_uri() + href.removeprefix("app.html")
    if href.startswith("index.html"):
        return (ROOT / "site" / "index.html").resolve().as_uri() + href.removeprefix("index.html")
    return href


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=50, start=100, bottom=50, end=100):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths[min(idx, len(widths) - 1)]
            cell.width = width
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_row_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        header = OxmlElement("w:tblHeader")
        header.set(qn("w:val"), "true")
        tr_pr.append(header)


def add_run_text(paragraph, text: str, bold=False, italic=False):
    run = paragraph.add_run(text)
    set_rfonts(run)
    run.bold = bold
    run.italic = italic
    return run


def clean_text(text: str) -> str:
    text = text.replace("\xa0", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def add_rich_text(paragraph, node: Tag):
    for child in node.children:
        if isinstance(child, NavigableString):
            txt = clean_text(str(child))
            if txt:
                add_run_text(paragraph, txt)
        elif isinstance(child, Tag):
            if child.name == "br":
                paragraph.add_run().add_break()
            elif child.name in ("strong", "b"):
                txt = clean_text(child.get_text(" ", strip=True))
                if txt:
                    add_run_text(paragraph, txt, bold=True)
            elif child.name in ("em", "i"):
                txt = clean_text(child.get_text(" ", strip=True))
                if txt:
                    add_run_text(paragraph, txt, italic=True)
            elif child.name == "code":
                txt = clean_text(child.get_text(" ", strip=True))
                if txt:
                    run = add_run_text(paragraph, txt)
                    run.font.name = "Consolas"
                    run.font.size = Pt(9.5)
            elif child.name == "a" and child.get("href"):
                txt = clean_text(child.get_text(" ", strip=True)) or child.get("href")
                href = normalize_href(child.get("href"))
                if href:
                    add_hyperlink(paragraph, txt, href)
                else:
                    add_run_text(paragraph, txt)
            else:
                txt = clean_text(child.get_text(" ", strip=True))
                if txt:
                    add_run_text(paragraph, txt)


def add_paragraph_from_node(doc, node: Tag, style=None):
    text = clean_text(node.get_text(" ", strip=True))
    if not text:
        return
    p = doc.add_paragraph(style=style)
    set_para_spacing(p, 0, 6, 1.10)
    add_rich_text(p, node)


def add_callout(doc, text: str):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    shade_cell(cell, CALLOUT_FILL)
    p = cell.paragraphs[0]
    set_para_spacing(p, 0, 0, 1.10)
    add_run_text(p, clean_text(text), bold=False)
    doc.add_paragraph()


def add_html_table(doc, html_table: Tag):
    rows = html_table.find_all("tr")
    if not rows:
        return
    max_cols = max(len(r.find_all(["th", "td"])) for r in rows)
    max_cols = max(max_cols, 1)
    widths = distribute_widths(max_cols)
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = "Table Grid"
    set_table_geometry(table, widths)

    for r_idx, html_row in enumerate(rows):
        set_row_cant_split(table.rows[r_idx])
        if r_idx == 0:
            set_row_repeat_header(table.rows[r_idx])
        cells = html_row.find_all(["th", "td"])
        for c_idx in range(max_cols):
            cell = table.cell(r_idx, c_idx)
            if r_idx == 0:
                shade_cell(cell, LIGHT_FILL)
            if c_idx < len(cells):
                text = clean_text(cells[c_idx].get_text(" ", strip=True))
            else:
                text = ""
            p = cell.paragraphs[0]
            set_para_spacing(p, 0, 0, 1.0)
            run = add_run_text(p, text, bold=(r_idx == 0 or cells[c_idx].name == "th" if c_idx < len(cells) else False))
            run.font.size = Pt(8.5 if max_cols >= 5 else 9)
            if r_idx == 0:
                run.font.color.rgb = INK
    doc.add_paragraph()


def distribute_widths(n: int):
    total = 9360
    if n == 1:
        return [total]
    if n == 2:
        return [2700, total - 2700]
    if n == 3:
        return [1800, 3660, 3900]
    if n == 4:
        return [1700, 2500, 2600, 2560]
    if n == 5:
        return [1300, 1800, 2300, 1800, 2160]
    if n == 6:
        return [1200, 1700, 2100, 1300, 1400, 1660]
    base = total // n
    widths = [base] * n
    widths[-1] += total - sum(widths)
    return widths


def process_node(doc, node: Tag, heading_offset=1):
    if isinstance(node, NavigableString):
        return
    if not isinstance(node, Tag):
        return
    if node.name in ("script", "style"):
        return

    if node.name in ("h1", "h2", "h3"):
        level = min(3, {"h1": 1, "h2": 2, "h3": 3}[node.name] + heading_offset)
        p = doc.add_paragraph(clean_text(node.get_text(" ", strip=True)), style=f"Heading {level}")
        return
    if node.name in ("p",):
        add_paragraph_from_node(doc, node)
        return
    if node.name == "blockquote":
        add_callout(doc, node.get_text(" ", strip=True))
        return
    if node.name == "pre":
        add_callout(doc, node.get_text("\n", strip=True))
        return
    if node.name in ("ul", "ol"):
        style = "List Bullet" if node.name == "ul" else "List Number"
        for li in node.find_all("li", recursive=False):
            p = doc.add_paragraph(style=style)
            set_para_spacing(p, 0, 4, 1.167)
            add_rich_text(p, li)
        return
    if node.name == "table":
        add_html_table(doc, node)
        return
    if node.name == "div" and "pipeline" in (node.get("class") or []):
        steps = node.find_all("div", class_="pl-step")
        if steps:
            for step in steps:
                p = doc.add_paragraph(style="List Number")
                set_para_spacing(p, 0, 4, 1.167)
                add_run_text(p, clean_text(step.get_text(" - ", strip=True)))
            return
    if node.name == "div" and "tiers" in (node.get("class") or []):
        tiers = node.find_all("div", class_="tier")
        if tiers:
            table = doc.add_table(rows=1, cols=4)
            set_table_geometry(table, [1700, 2100, 2860, 2700])
            headers = ["플랜", "가격", "포함 범위", "초과/비고"]
            for idx, header in enumerate(headers):
                shade_cell(table.cell(0, idx), LIGHT_FILL)
                p = table.cell(0, idx).paragraphs[0]
                add_run_text(p, header, bold=True)
            for tier in tiers:
                row = table.add_row().cells
                parts = tier.find_all("div")
                texts = [clean_text(p.get_text(" ", strip=True)) for p in parts]
                padded = (texts + ["", "", "", ""])[:4]
                for idx, text in enumerate(padded):
                    p = row[idx].paragraphs[0]
                    run = add_run_text(p, text)
                    run.font.size = Pt(9)
            doc.add_paragraph()
            return

    direct_table = node.find("table", recursive=False)
    if direct_table and len(list(node.children)) <= 3:
        for child in node.children:
            process_node(doc, child, heading_offset)
        return

    text = clean_text(node.get_text(" ", strip=True))
    if node.name == "div" and text and len(text) < 300 and not node.find(["p", "ul", "ol", "table", "h1", "h2", "h3"], recursive=False):
        p = doc.add_paragraph()
        set_para_spacing(p, 0, 4, 1.10)
        add_run_text(p, text)
        return

    for child in node.children:
        process_node(doc, child, heading_offset)


def page_html_map():
    soup = BeautifulSoup(INDEX.read_text(encoding="utf-8"), "html.parser")
    return {script["data-page"]: script.decode_contents() for script in soup.select("script[data-page]")}


def nav_page_titles():
    soup = BeautifulSoup(INDEX.read_text(encoding="utf-8"), "html.parser")
    node = soup.select_one("#nav-data")
    if not node:
        return {}
    data = json.loads(node.get_text())
    out = {}
    for book in data.get("books", []):
        for group in book.get("groups", []):
            for page in group.get("pages", []):
                out[page["id"]] = page["title"]
    return out


def add_cover(doc: Document):
    p = doc.add_paragraph()
    set_para_spacing(p, 0, 4, 1.0)
    run = add_run_text(p, "AX Product Manager (Ontology) 과제")
    run.font.size = Pt(23)
    run.font.bold = True
    run.font.color.rgb = INK

    p = doc.add_paragraph()
    set_para_spacing(p, 0, 16, 1.15)
    run = add_run_text(p, "Ontology 기반 재무·회계 AI 자동화 솔루션 제품화 전략")
    run.font.size = Pt(13)
    run.font.color.rgb = MUTED

    rows = [
        ("대상", "뤼튼테크놀로지스 AX Product Manager (Ontology) 과제 전형"),
        ("문서 성격", "제출용 Word 초안"),
        ("작성일", date.today().isoformat()),
        ("화면 설계", "통합 인터랙티브 데모 v1.0 및 화면 설계 기획서 반영"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    set_table_geometry(table, [1600, 7760])
    table.style = "Table Grid"
    for i, (label, value) in enumerate(rows):
        shade_cell(table.cell(i, 0), LIGHT_FILL)
        p0 = table.cell(i, 0).paragraphs[0]
        p1 = table.cell(i, 1).paragraphs[0]
        add_run_text(p0, label, bold=True)
        add_run_text(p1, value)

    doc.add_paragraph()
    add_callout(
        doc,
        "본 초안은 PRD, 화면 설계, 제품 로드맵, 가격 전략을 하나의 Word 제출본으로 묶은 문서입니다. 화면 설계 산출물은 통합 데모와 상세 기획서 내용을 함께 반영했습니다.",
    )


def add_document_index(doc: Document, titles: dict[str, str]):
    doc.add_paragraph("문서 인덱스", style="Heading 1")
    p = doc.add_paragraph()
    set_para_spacing(p, 0, 8, 1.10)
    add_run_text(p, "아래 인덱스는 본 Word 초안에 포함된 전체 파트와 하위 섹션을 기준으로 정리한 정적 목차입니다. ")
    add_run_text(p, "각 파트의 번호는 본문 Heading 구조와 동일하게 맞췄습니다.", bold=True)

    groups = [
        ("1. 제품 기획서 (PRD)", [f"prd-{i}" for i in range(1, 20)]),
        ("2. 화면 설계", ["ui-intro", "ui-admin", "ui-reviewer", "ui-onboarding"]),
        ("3. 제품 로드맵", [f"rm-{i}" for i in range(1, 7)]),
        ("4. 가격 전략", [f"pr-{i}" for i in range(7, 13)]),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_geometry(table, [2200, 3560, 3600])
    headers = ["파트", "섹션", "주요 내용"]
    for i, h in enumerate(headers):
        shade_cell(table.cell(0, i), LIGHT_FILL)
        add_run_text(table.cell(0, i).paragraphs[0], h, bold=True)
    summaries = {
        "1. 제품 기획서 (PRD)": "제품 비전, 문제 정의, MVP 워크플로우, 기능 요구사항, 성공 지표, 리스크",
        "2. 화면 설계": "통합 데모 v1.0, 관리자 8탭, 검토자 2개 워크플로우, 온보딩 7단계",
        "3. 제품 로드맵": "Phase별 배포 계획, 4인 스쿼드 타임라인, 의존성, 리스크, 졸업 게이트",
        "4. 가격 전략": "피부과·성형외과 버티컬 가격, 파트너 채널, 성과 과금, 경쟁 포지셔닝",
    }
    for part, page_ids in groups:
        cells = table.add_row().cells
        add_run_text(cells[0].paragraphs[0], part, bold=True)
        add_run_text(cells[1].paragraphs[0], " / ".join(titles.get(pid, pid) for pid in page_ids))
        add_run_text(cells[2].paragraphs[0], summaries[part])
    doc.add_paragraph()


def add_screen_design_section(doc: Document, html_map):
    doc.add_paragraph("2. 화면 설계", style="Heading 1")
    app_uri = (ROOT / "site" / "app.html").resolve().as_uri()
    p = doc.add_paragraph()
    set_para_spacing(p, 0, 6, 1.10)
    add_run_text(p, "화면 설계는 통합 인터랙티브 데모와 웹 문서의 화면 설계 기획서 내용을 함께 반영했습니다. ")
    add_hyperlink(p, "통합 데모 열기", f"{app_uri}?screen=admin")
    doc.add_paragraph()

    for page_id in ["ui-intro", "ui-admin", "ui-reviewer", "ui-onboarding"]:
        html = html_map.get(page_id)
        if not html:
            continue
        soup = BeautifulSoup(html, "html.parser")
        for child in soup.contents:
            process_node(doc, child, heading_offset=1)


def add_pages(doc, pages, heading: str, html_map, heading_offset=1):
    doc.add_paragraph(heading, style="Heading 1")
    for page_id in pages:
        html = html_map.get(page_id)
        if not html:
            continue
        soup = BeautifulSoup(html, "html.parser")
        for child in soup.contents:
            process_node(doc, child, heading_offset=heading_offset)


def main():
    html_map = page_html_map()
    titles = nav_page_titles()
    doc = Document()
    style_document(doc)
    add_cover(doc)
    doc.add_page_break()
    add_document_index(doc, titles)
    doc.add_page_break()

    prd_pages = [f"prd-{i}" for i in range(1, 20)]
    add_pages(doc, prd_pages, "1. 제품 기획서 (PRD)", html_map)

    doc.add_page_break()
    add_screen_design_section(doc, html_map)

    doc.add_page_break()
    roadmap_pages = [f"rm-{i}" for i in range(1, 7)]
    add_pages(doc, roadmap_pages, "3. 제품 로드맵", html_map)

    pricing_pages = [f"pr-{i}" for i in range(7, 13)]
    add_pages(doc, pricing_pages, "4. 가격 전략", html_map)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
